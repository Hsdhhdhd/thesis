from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("thesis_中文逐节全对照阅读版.docx")


def set_font(run, size=10, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, size=9.3, bold=False):
    cell.text = ""
    for idx, part in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(part)
        set_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, size=14 if level == 1 else 11.5, bold=True)


def add_pair(doc, left, right):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for i, width in enumerate((Cm(11.8), Cm(14.8))):
        table.columns[i].width = width
        table.rows[0].cells[i].width = width
    set_cell_text(table.rows[0].cells[0], left, size=8.8)
    set_cell_text(table.rows[0].cells[1], right, size=9.5)


def add_header_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    labels = ("Original location / 英文原文位置", "Chinese counterpart / 中文对应阅读版")
    for cell, label in zip(table.rows[0].cells, labels):
        shade(cell, "D9EAF7")
        set_cell_text(cell, label, size=10, bold=True)


CONTENT = [
    (
        "Front Matter",
        "Abstract",
        "本论文开发并评估了一套可穿戴惯性测量单元（IMU）触觉反馈系统，用于监测长时间坐姿中的骨盆旋转。原型由佩戴在骶骨附近的 Seeed Studio XIAO nRF54L15 Sense、DRV2605L 触觉驱动器和硬币式振动电机组成。Android 应用通过低功耗蓝牙记录带标签的 IMU 数据，PC 工具用于信号检查。在与 Shimmer 参考设备的趋势级比较中，XIAO 在共位测试中的 Pearson 相关系数达到 0.816-0.923；现场数据经过滞后校正后，最大互相关最高达到 0.981。\n\n五名参与者的初步现场研究收集了列车旅途中按指令标注的坐姿片段。在单一留出参与者划分中，k=7 的 K 近邻（KNN）将 2 秒窗口分类为静态坐姿或骨盆旋转，F1 分数为 0.957；165 个骨盆旋转窗口中检测到 155 个，3,497 个静态坐姿窗口中只有 4 个误报。一个可部署的紧凑原型 KNN（12 个特征，128 个 KMeans 原型）在同一留出参与者上达到 F1=0.959，并已部署到可穿戴设备上。\n\n每个 2 秒分类周期在设备端需要 2.73 毫秒，相当于低于窗口周期 0.15% 的计算时间，固件占用 56.3 KB flash 和 9.3 KB RAM。固件实现了 10 分钟无活动计时器，触发 1 秒触觉提醒。这些结果支持该指令式初步任务的技术可行性，但尚不能证明自然日常使用中的鲁棒性、电池寿命或提醒的行为效果。"
    ),
    (
        "Front Matter",
        "Keywords",
        "关键词：嵌入式机器学习；触觉反馈；惯性测量单元；骨盆旋转；长时间坐姿；可穿戴感知。"
    ),
    (
        "Chapter 1: Introduction",
        "1.1 Background and Motivation - sedentary behaviour and operational target",
        "久坐行为通常定义为清醒状态下、以坐、斜靠或躺卧姿势进行、能量消耗不超过 1.5 个代谢当量（MET）的行为。标准能量消耗参考中，坐姿本身通常处于 1.0-1.5 MET 范围内；久坐生活方式研究也强调，这类低能量行为不能简单等同于“缺乏运动”。在这个大背景下，本文聚焦于更窄的操作目标：长时间静态坐姿，即坐姿期间下躯干和骨盆活动受限的时段。原型并不估算能量消耗，也不分类所有久坐行为，而是监测坐姿佩戴者在较长坐姿时段内是否发生了可检测的骨盆运动。"
    ),
    (
        "Chapter 1: Introduction",
        "1.1 Background and Motivation - why static sitting differs from sedentary time",
        "这个区分很重要，因为久坐时间和静态坐姿相关，但并不相同。一个人可以仍然处于低能量的久坐状态，同时改变姿势或移动骨盆。对于需要触发局部反馈的可穿戴系统而言，必须有一个可检测的运动目标。工作成年人可能积累大量职业坐姿时间，人口数据也显示了工作相关体力活动的长期变化。关于久坐中断的证据表明，久坐如何积累也很重要。因此，本文真正面对的设计问题不是简单判断一个人是否坐着，而是判断一个较长的坐姿时段是否在骨盆周围保持静止。"
    ),
    (
        "Chapter 1: Introduction",
        "1.1 Background and Motivation - health context and scope boundary",
        "长时间坐姿与心代谢风险因素、心血管疾病、2 型糖尿病、肥胖和全因死亡率有关。坐姿也会影响脊柱和腰骶区域的机械负荷，坐姿行为还被用于研究久坐办公室工作者的下背痛。不过，本文不诊断姿势、不估算脊柱负荷，也不提出关于背痛的临床主张。这些文献主要用于支持一个实用的可穿戴提示系统：帮助打断长时间静态坐姿。站立休息和短时活动提示可以减少坐姿时间，主动工作站也被评估过对工作效率和表现的影响，但这些方法并不是骨盆局部感知和提示传递系统。佩戴式系统可以通过感知相关身体节段并在身体上直接发出提示来填补这个技术空白。"
    ),
    (
        "Chapter 1: Introduction",
        "1.2 Problem Statement",
        "本文要解决的具体问题是：目前缺少一种小型、移动、聚焦骨盆的反馈系统，能够检测静态坐姿并触发局部触觉提示，以打断长时间静态坐姿。现有监测方法可以描述坐姿持续时间、总体活动状态或坐姿姿态，但它们本身并不能完整实现从骨盆层面运动感知到即时局部反馈的技术链。该系统应能在不同座椅和日常情境中工作，不依赖摄像头或传感家具，也不要求持续把原始数据传输到手机或云端服务。同时，它必须足够小，能够附着在身体上，并且足够简单，能在低功耗设备上嵌入式执行。"
    ),
    (
        "Chapter 1: Introduction",
        "1.2 Problem Statement - three linked parts",
        "这个问题包含三个相互关联的部分。第一，传感位置必须匹配目标运动。手腕或手机传感器可以检测许多日常活动，但也可能对打字、用手机或手臂动作产生反应，而此时骨盆仍然静止。第二，系统必须能及时反馈。只能在会话结束后分析的数据记录器对研究有用，但不能在坐姿过程中提示佩戴者。第三，系统必须具备技术部署可行性。依赖笔记本电脑、云服务或大型存储数据集的分类器，不符合本地触觉反馈的目标。"
    ),
    (
        "Chapter 1: Introduction",
        "1.2 Problem Statement - comparison with existing approaches",
        "基于视觉的系统可以在固定环境中分类坐姿，带传感器的椅子可以通过压力模式检测坐姿活动。这些方法与本文相关，但它们最适合受控环境，不能跟随用户跨越不同日常场景。腕戴式可穿戴设备解决了移动性问题，但对本文任务来说测量的是更间接的身体节段：智能手表可能记录到打字或手机使用，而骨盆仍然静止。Shimmer 等研究级惯性传感器适合作为参考测量，但更偏向离线数据采集，而不是低成本本地反馈。"
    ),
    (
        "Chapter 1: Introduction",
        "1.2 Problem Statement - proposed system labels",
        "因此，本文系统把传感器放在骨盆附近，并把反馈保持在本地。目标运动称为“骨盆旋转”：由骨盆安装 IMU 检测到的、有意的坐姿前后向骨盆运动。对照类别是“静态坐姿”：没有有意骨盆运动的坐姿时段。在本文中，这些标签是技术检测类别，不是临床姿势类别。"
    ),
    (
        "Chapter 1: Introduction",
        "1.3.1 Sensing Modalities for Seated Movement",
        "检测坐姿状态、姿态或运动的系统大致可分为三类：摄像头、带传感器的家具和身体佩戴式传感器。视觉方法从 RGB、深度或 RGB-D 数据估计人体姿态；在坐姿场景中，机器视觉也被用于分类正确和不正确坐姿。它们的优势是空间信息丰富，但固定在特定房间，依赖视线，并带来隐私问题。传感椅和智能坐垫通过压力或接触变化分类坐姿。它们避免了摄像头隐私问题，但把测量绑定到单一座位，不适合公共交通、家庭、办公室和共享工作空间。"
    ),
    (
        "Chapter 1: Introduction",
        "1.3.1 Sensing Modalities - wearable focus",
        "身体佩戴式感知可以跟随用户跨环境移动。对本文而言，这种移动性比重建完整坐姿更重要。相关任务是在短时间序列窗口中识别是否发生了骨盆运动。这个选择有意缩小了感知目标：设备并不判断姿势是否符合人体工学，而是判断佩戴者是否已经在骨盆周围保持静止足够久，从而需要触发提醒。"
    ),
    (
        "Chapter 1: Introduction",
        "1.3.2 Wearable IMU-Based HAR",
        "IMU 结合加速度计和陀螺仪，体积小，可以在非准备环境中使用，因此适合可穿戴运动识别和脊柱相关运动评估。传感器位置是核心设计选择：对于坐姿骨盆运动，骨盆和下躯干比手腕或大腿更直接提供信息。骶骨或下背部位置也实际可行，因为它接近目标运动，同时不占用双臂。"
    ),
    (
        "Chapter 1: Introduction",
        "1.3.2 Wearable IMU-Based HAR - feedback gap",
        "可穿戴生物反馈研究表明，身体佩戴式感知可以支持运动反馈，但很多系统仍然以研究为导向并依赖离线分析。在这种情况下，传感器可以描述发生了什么，但不能自主决定何时提示佩戴者。这为低成本嵌入式可穿戴设备留下空间：它既保留相关骨盆运动信息，又能实现即时本地反馈。"
    ),
    (
        "Chapter 1: Introduction",
        "1.3.3 Machine Learning and TinyML",
        "基于传感器的人体活动识别越来越多使用深度学习，但在小型可穿戴微控制器上，传统方法仍然有吸引力。支持向量机、决策树、随机森林和 KNN 都是运动识别中的相关传统基线。连续数据通常被划分为固定长度窗口，再用均值、标准差、范围、均方根等轻量时域特征概括。低能耗 HAR 硬件研究也说明，在可穿戴部署中，特征和模型的简单性很重要。本文采用窗口方法是合适的，因为提醒逻辑不需要连续解剖角度估计，只需要重复判断最近窗口是否包含骨盆运动。"
    ),
    (
        "Chapter 1: Introduction",
        "1.3.3 TinyML and compact KNN",
        "对于闭环可穿戴反馈，推理发生在哪里非常关键。边缘 AI 和 TinyML 把推理移到设备本身，但受到内存、计算时间和能耗限制。KNN 简单、可解释，适合作为基线，但通常要存储所有训练样本。凝聚近邻方法通过用更小的代表性子集替代完整参考集来解决存储问题。本文使用 KMeans 聚类作为实际的原型生成步骤，用少量类别条件原型替代每个类别的样本集合。"
    ),
    (
        "Chapter 1: Introduction",
        "1.3.4 Research Gap",
        "本文由三个空白推动。第一，很多坐姿监测系统关注坐姿时间、宽泛姿态类别或固定椅子使用，而本文目标是更窄的坐姿内部运动：最近一段时间骨盆是否移动。第二，许多可穿戴 IMU 研究作为离线感知或分类管线评估，而提醒系统还需要本地决策和提示传递。第三，面向现场使用的低成本原型需要在链条多个环节提供证据：与参考传感器的一致性、可用的带标签数据采集、跨人检测性能和嵌入式计时。"
    ),
    (
        "Chapter 1: Introduction",
        "1.3.4 Research Gap - thesis response",
        "本文通过实现一条完整工作管线来回应这些空白：低成本骨盆区域 IMU 可穿戴设备采集带标签现场数据，与 Shimmer 参考进行检查，在设备端分类静态坐姿和骨盆旋转，并提供触觉反馈。视觉和智能家具系统不能跟随用户；腕戴设备对骨盆运动测量较间接；研究级 IMU 通常用于事后分析而不是自主反馈。因此，本文并不声称解决所有久坐行为监测问题，而是测试用于本地打断长时间静态坐姿所需的技术链能否在小型可穿戴设备上运行。"
    ),
    (
        "Chapter 1: Introduction",
        "1.3.4 Novelty and scope",
        "本文的新意不是提出新的机器学习算法，也不是提出临床姿势指标，而是集成并初步验证从感知到反馈的完整技术路径。这个范围也限定了工作的边界：研究可以评估传感器一致性、数据采集、分类和嵌入式计时，但尚不能证明行为改变或健康收益。"
    ),
    (
        "Chapter 1: Introduction",
        "1.4 Research Aim and Questions",
        "本文目标是设计、实现并初步评估一种基于 IMU 的可穿戴触觉反馈系统，用于监测长时间坐姿中的骨盆运动。研究由三个问题引导：第一，佩戴在骨盆附近的低成本嵌入式 IMU 能否捕捉到与研究级 Shimmer IMU 足够一致的运动模式？第二，所提出的可穿戴系统能否在真实现场和交通情境中支持实时数据采集与标注？第三，采集到的 IMU 数据能否区分静态坐姿和骨盆旋转，并且紧凑分类器能否高效部署到设备端，以支持打断长时间静态坐姿的触觉反馈？"
    ),
    (
        "Chapter 1: Introduction",
        "1.5 Contributions and 1.6 Thesis Outline",
        "主要贡献包括：基于 XIAO nRF54L15 Sense、DRV2605L 驱动器和硬币式振动电机的紧凑骨盆区域可穿戴原型；带多设备 BLE 采集、实时可视化和标签记录的 Android 与 PC 数据采集平台；基于 Shimmer 的传感器比较协议；包含步行、公共交通和坐姿列车旅途片段的五名参与者带标签初步现场数据集；以及端到端设备端机器学习管线，包括窗口特征提取、五种分类器比较、KMeans 原型压缩、嵌入式部署和设备端推理计时。\n\n论文结构为：第 2 章介绍硬件和传感器位置；第 3 章介绍 Android/PC 数据采集工具和设备端分类管线；第 4 章报告 Shimmer 验证和初步检测研究；第 5 章讨论结果与局限；第 6 章总结全文。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.1 System Overview",
        "该原型是一个身体佩戴单元，用于在长时间坐姿中监测骨盆旋转，并在设定时间内未检测到骨盆旋转时提供触觉反馈。设计直接来自第 1 章提出的研究空白：系统必须感知正确的身体节段，支持带标签现场记录，在本地运行最终决策，并提供即时提示。因此设备必须佩戴在骨盆附近，因为目标运动发生在下躯干而不是手腕；必须在现场会话中记录带标签数据，因为分类器需要真实坐姿场景中的标注记录；必须本地运行最终决策，因为触觉反馈不应依赖笔记本或云连接；同时还要保持低成本和足够紧凑，便于反复原型开发。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.1 Functional architecture",
        "系统功能架构如下：XIAO nRF54L15 Sense 上的 IMU 在骨盆附近测量三轴加速度和角速度；微控制器处理这些值并追踪是否发生骨盆旋转。如果在固件设定的 10 分钟间隔内没有检测到骨盆旋转，微控制器驱动 DRV2605L 触觉驱动器，使硬币式振动电机振动 1 秒。同一数据也可以通过 BLE 传输到智能手机或 PC，用于记录、标注和离线分析。研究流程和最终反馈环路是分开的：采集阶段用 BLE 流检查信号和存储标签文件；部署阶段紧凑分类器和无活动计时器在可穿戴设备上运行。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.2 Mechanical Prototype",
        "原型由 3D 打印的多孔载板和橙色外壳组装而成。XIAO nRF54L15 Sense 与 DRV2605L 安装在黑色载板上，并用热熔胶机械固定。500 mAh 电池放在外壳内部，外壳尺寸为 5.4 × 3.0 × 1.5 cm。原型部件成本低于 30 欧元，不包括手机、PC 和 Shimmer 参考设备。外壳被设计为功能性研究原型，而不是最终消费产品：打印主体保护电池并提供平整接触面，多孔载板让电子部件可见，便于调试。这种选择方便开发阶段调整接线和固件，但也限制了机械鲁棒性。后续现场使用中，电池接触方式被发现是间歇性记录缺口的来源之一。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.2 Main hardware components",
        "主要硬件部件包括：XIAO nRF54L15 Sense 作为微控制器、嵌入式 IMU 和 BLE 无线模块；LSM6DS3TR-C IMU 提供三轴加速度和角速度；DRV2605L 驱动振动电机；硬币式振动电机向佩戴者提供触觉提醒；500 mAh 电池作为内部电源；3D 打印外壳用于机械封装。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.3 Electrical Integration",
        "XIAO 和 DRV2605L 通过电源线和 I2C 总线连接。XIAO 的 GND 和 3V3 分别连接 DRV2605L 的 GND 和 VIN；XIAO D4 作为串行数据线 SDA，D5 作为串行时钟线 SCL。振动电机连接到 DRV2605L 的正负电机输出端。接线有意保持最小化：IMU 已集成在 XIAO 板上，外部电路只增加触觉驱动器、电机和电池。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.4 Sensor Node and Configuration",
        "选择 Seeed Studio XIAO nRF54L15 Sense，是因为它把微控制器、LSM6DS3TR-C IMU 和 BLE 无线模块集成在一个小型板上。与分离式微控制器和传感器模块相比，这减少了接线、板面积和部件数量。加速度计和陀螺仪配置为 208 Hz，这为短时骨盆旋转运动提供足够时间分辨率，同时仍适合无线记录。传感器流包含六个原始通道：三轴加速度 ax、ay、az 和三轴角速度 gx、gy、gz。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.4 Simple IMU configuration",
        "IMU 配置故意保持简单。论文使用原始加速度和角速度通道，然后在离线处理和嵌入式推理期间提取基于幅值的特征。这样可以避免依赖绝对姿态估计，因为后者需要额外校准，并且对具体佩戴位置更敏感。对于本文的二分类决策，最近窗口内是否存在运动，比重建精确骨盆角度更重要。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.5 Haptic Feedback Module",
        "原型中的电机是 8 mm × 3 mm 硬币式振动电机，供应商给出的额定电压为 3-5 V、额定电流为 67 mA。由于未找到该具体零售件的正式数据表，论文还参考了 Precision Microdrives 的相近 8 mm 硬币振动电机数据表，其给出 8 mm 直径、3.4 mm 机身长度、3 V 额定工作电压和 66 mA 典型额定电流。该电机足够小，可以放入外壳，并在下背部提供局部触觉输出。使用 DRV2605L 是为了避免微控制器直接驱动电机，并为未来支持更可控的触觉模式留出空间。当前固件中，无活动阈值为 10 分钟，满足条件时电机振动 1 秒；振动强度固定，后续可根据用户偏好调整模式、强度或间隔。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.5 Haptic scope",
        "触觉输出在本文中是提示传递机制，不是已经验证的行为干预。它的作用是说明感知和分类结果可以连接到即时本地提醒。振动在长期使用中是否明显、可接受或有效，是未来工作，而不是当前设计可以默认成立的结论。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.6 Wearing Position",
        "设备佩戴在接近骶骨的下背部/骨盆区域。这个位置比手腕更直接反映下躯干和骨盆运动，并且与 Liao 等人针对 IMU 骨盆旋转检测所报告的骶骨/骨盆放置一致。该位置也接近振动目标，因此同一身体位置同时用于感知和反馈。"
    ),
    (
        "Chapter 2: Wearable System Design",
        "2.6 Body attachment",
        "身体固定方式为：外壳用双面胶固定到 10 cm × 15 cm 的无菌敷贴上，再将敷贴贴到皮肤上。这种两层结构分散胶粘负荷，并允许每次会话之间更换敷贴而不重新定位设备。它适合短期初步测试，因为简单、便宜并能适配不同参与者，但不应视为长期日常佩戴的最终固定方案。"
    ),
    (
        "Chapter 3: Implementation",
        "Chapter opening",
        "第 3 章描述用于采集带标签 IMU 数据的软件平台、离线分类器比较以及运行在 XIAO nRF54L15 Sense 上的紧凑 KNN 模型。实现围绕同一条数据管线设计：可穿戴设备采样 IMU，在实验期间把同一六通道数据流传输到采集工具，之后在设备端运行训练后检测器的简化版本。让不同工具共享同一数据包格式，可以降低台架测试、现场采集和嵌入式部署描述不同信号的风险。\n\n系统开发了两个用于同一工作流不同阶段的采集工具：Android 应用用于移动现场采集，PC 接收器用于台架测试、调试和可视化。"
    ),
    (
        "Chapter 3: Implementation",
        "3.1 BLE Data Interface",
        "Android 应用和 PC 工具使用同一个 BLE 接口。可穿戴设备广播名为 ble-imu-sensor；接收端订阅固件中定义的自定义服务和通知特征。每条通知携带 12 字节载荷，解码为六个小端有符号 16 位整数，对应三轴加速度和三轴角速度。在 BLE 流固件中，加速度以 10^-3 m/s² 编码，角速度以毫弧度每秒编码。因此两个接收器都会把加速度除以 1000 得到 m/s²，并把陀螺仪值乘以 180/(π·10^3) 约等于 0.0573，得到度每秒。"
    ),
    (
        "Chapter 3: Implementation",
        "3.1 Payload rationale",
        "载荷有意设计得很紧凑。12 字节通知只包含六个运动值，记录时的时间戳由接收设备负责。这样可以保持无线传输简单，避免更大的数据包格式，同时保留分类器需要的所有通道。嵌入式推理期间，固件通过 Zephyr 传感器 API 读取 m/s² 加速度和 rad/s 角速度，再把角速度转换为度每秒，然后进行特征提取，从而让部署模型与记录训练数据使用同一单位约定。"
    ),
    (
        "Chapter 3: Implementation",
        "3.2 Android Field Collection",
        "Android 应用扫描附近可穿戴设备，通过 BLE 连接，订阅 IMU 通知特征，并记录每个已连接设备的最新加速度和陀螺仪值。它可以同时连接两个可穿戴设备，这对多设备现场会话有用。记录时，应用把数据包写入手机外部文档目录中的 CSV 文件。前台服务在锁屏时保持 BLE 和 GPS 采集。GPS 样本来自 Android 融合定位提供器；每行 IMU 数据都包含 gps_age_ms，即 IMU 包和最近一次 GPS 定位之间的时间差。"
    ),
    (
        "Chapter 3: Implementation",
        "3.2 Android annotation design",
        "应用支持按设备标注活动。记录开始后，每个连接的可穿戴设备都有独立标签控件；所选标签会写入该设备收到的每一行 CSV，直到标签停止。默认标签可配置，因此研究者可以匹配协议，例如静态坐姿、骨盆旋转、步行或交通片段。每个标注分配一个片段 ID 和开始时间戳，使后续处理能够恢复带标签时间片段，而不是把每行当作孤立样本。"
    ),
    (
        "Chapter 3: Implementation",
        "3.2 Why row-level labels matter",
        "这种标注设计对现场协议很重要。标签不是由手机自动推断，而是在会话期间输入，之后用于筛选干净的坐姿片段。把标签写入每一行便于处理文件，而片段 ID 保留连续行属于同一指令活动阶段这一事实。没有活动标注的行可以保留在原始日志中以保证完整性，但会从分类器数据集中排除。"
    ),
    (
        "Chapter 3: Implementation",
        "3.2 Android CSV fields",
        "Android CSV 每个数据包一行，包含设备身份、手机接收时间戳、六个转换后的 IMU 通道、GPS 元数据和标注字段。没有活动标注的行在标注列中留空。字段组包括设备和时间、IMU 值、GPS 时间、GPS 位置、GPS 质量、GPS 运动状态和活动标注。"
    ),
    (
        "Chapter 3: Implementation",
        "3.3 PC Receiver and Visualisation",
        "PC 接收器用 Python 实现，BLE 使用 bleak，界面使用 PyQt。它使用与 Android 应用相同的通知特征，但角色不同：用于台架测试、调试和实时信号检查。界面显示最新加速度、去重力运动加速度、角速度和四元数估计，并提供原始加速度、角速度和运动加速度的实时曲线。为避免阻塞 GUI，收到的数据包进入队列，由专门线程写入 CSV。接收器还可以把缓存数据导出为 Excel 文件。"
    ),
    (
        "Chapter 3: Implementation",
        "3.3 PC role and CSV fields",
        "PC 工具主要用于在现场会话前后检查原型是否正确响应。实时图可以立即显示传感器饱和、设备断连或异常轴向行为。PC 日志中的额外派生通道不是 Android 现场数据集所必需的，但在开发过程中有用，因为它们让运动峰值更容易检查。与 Android 日志相比，PC CSV 省略标注和 GPS 字段，但增加去重力运动加速度和四元数估计，便于台架测试中的实时信号检查。"
    ),
    (
        "Chapter 3: Implementation",
        "3.4 Windowing and Feature Extraction - dataset cleaning",
        "对于离线分类器，列车旅途中带标签坐姿片段经过人工筛选，以确保包含连续 IMU 流且没有电池断连缺口。这个选择步骤把分类实验聚焦在运动类别上，而不是缺失数据伪影上。保留的数据集包含 98 个导出片段文件，其中 66 个为静态坐姿片段，32 个为骨盆旋转片段。窗口化后得到 12,489 个静态坐姿窗口和 564 个骨盆旋转窗口。活动标签之外的行、非坐姿路线部分以及受可见电池接触缺口影响的片段在特征提取前被排除。加载过程中，非数值样本被删除，行按 epoch 时间戳排序，并移除完全重复行。"
    ),
    (
        "Chapter 3: Implementation",
        "3.4 Windowing",
        "在这些干净片段中，原始 IMU 数据被划分为 2 秒滑动窗口，步长为 1 秒，只保留至少包含 50 个样本的窗口。每个保留窗口继承其来源片段的标签。IMU 配置为 208 Hz，但轮询固件中的 I2C 开销和 RTOS 调度延迟使干净片段中的有效读取频率约为 172 Hz，因此每个 2 秒窗口大约包含 344 个样本。"
    ),
    (
        "Chapter 3: Implementation",
        "3.4 Feature extraction",
        "三个派生信号为加速度幅值、去重力加速度幅值以及陀螺仪幅值。对于六个原始通道和这三个派生幅值信号，分别计算十个时域统计量：均值、标准差、最小值、最大值、范围、均方根、中位数、四分位距、一阶差分平均绝对值和一阶差分标准差。因此，每个窗口得到 90 维特征向量。"
    ),
    (
        "Chapter 3: Implementation",
        "3.5 Compact KNN for Embedded Deployment",
        "存储全部 9,391 个训练窗口的完整 KNN 基线不适合低功耗可穿戴设备。部署模型采用两项压缩：人工选择 12 个特征，以及 KMeans 原型压缩。该特征子集在最终参与者 5 测试前固定，并基于参与者 1-4 的信号可解释性、嵌入式实现成本和内部训练折行为选取。选中特征主要来自陀螺仪和加速度范围信息，这与静态坐姿和有意骨盆运动之间预期的信号差异一致。参与者 5 没有被用于选择特征子集、缩放值、原型数量或 KNN 参数。"
    ),
    (
        "Chapter 3: Implementation",
        "3.5 Scaling, prototypes, and KNN",
        "压缩和推理前，每个选中特征用只在参与者 1-4 上拟合的均值和标准差进行标准化；这些缩放值与原型集一起导出到固件头文件中。在标准化后的训练窗口中，每个标签类别内使用 KMeans 提取 64 个中心，因此总共存储 128 个原型。推理时，当前窗口与这些原型比较，使用 k=7 的距离加权 KNN。原型数量是在四名训练参与者上内部留一参与者验证后的内存/性能折中：每类 64 个原型相较 32 个提高平均内部 F1，同时保持参考集紧凑；128 个原型只带来很小额外收益但存储成本翻倍。"
    ),
    (
        "Chapter 3: Implementation",
        "3.5 Selected compact features",
        "紧凑 KNN 的 12 个特征包括：gyro_x_rms、gyro_x_std、gyro_x_range；gyro_mag_std、gyro_mag_max、gyro_mag_rms、gyro_mag_range、gyro_mag_mean；accel_z_max、accel_z_range、accel_z_std；以及 acc_motion_mag_std。它们集中反映角速度变化、角速度幅值和 z 轴加速度变化。"
    ),
    (
        "Chapter 3: Implementation",
        "3.5 Firmware integration",
        "紧凑 KNN 集成到固件中。设备以 5 ms 轮询间隔采样 IMU，累积非重叠 2 秒窗口，在线计算 12 个窗口级特征，并将每个窗口分类为静态坐姿或骨盆旋转。嵌入式周期使用与离线分析相同的 2 秒窗口长度和特征定义，但不使用离线 1 秒重叠步长。当检测到骨盆旋转时，无活动计时器重置；否则，如果 10 分钟内没有检测到骨盆旋转窗口，DRV2605L 驱动振动电机 1 秒。设备端每窗口推理延迟用硬件周期计数器测量。部署时，默认 Zephyr 主线程栈在同时使用 I2C 驱动、日志和窗口处理时过小；把主线程栈设为 4,096 字节后解决了启动时的静默栈溢出故障。"
    ),
    (
        "Chapter 4: Evaluation",
        "Chapter opening",
        "评估分为两部分。第一，将嵌入式 IMU 与 Shimmer 参考传感器比较，以检查自制设备是否捕捉到可比的运动趋势。第二，使用五名志愿者的初步现场研究，测试 2 秒 IMU 窗口能否被分类为静态坐姿或骨盆旋转，并测量紧凑部署模型的性能和计时。本章报告方法和结果，结果含义的解释留到第 5 章讨论。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.1.1 Shimmer setup and data",
        "在共位比较中，自制可穿戴设备和 Shimmer 参考设备固定在同一粘贴安装面上，并作为一个刚体一起移动。Shimmer 配置为 100 Hz，XIAO IMU 配置为 208 Hz。两者被手动进行多次旋转、摆动和静止阶段。原计划为五名参与者收集配对 Shimmer/XIAO 记录，但只有两名产生了可用配对记录；另外三名的 Shimmer 数据未成功记录。这些失败的参考日志没有用于 Shimmer 比较，但对应 XIAO 记录中可用的干净带标签坐姿片段仍用于分类器数据集。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.1.1 Long field-worn recordings",
        "在每个有效配对案例中，两个传感器都在城市长路线中佩戴在骨盆附近，路线包括从 Leuven Group T 出发步行、列车旅行、用餐、返程和步行返回。这些长记录只用于比较设备之间的趋势一致性；分类器数据集使用后文筛选的带标签坐姿片段。XIAO 日志中 P1 有 57 个超过 1 秒的缺口，P2 有 4 个缺口。表中总采样率包含这些死时间；用于分类器的干净带标签片段约为 172 Hz。这些缺口的硬件含义在讨论章说明。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.1.2 Synchronisation and Metrics",
        "两个系统使用独立时钟。对于现场数据，epoch 时间戳定义重叠区间；对于共位记录，信号被重采样到共同时间基，并通过互相关对齐。现场信号重采样到 10 Hz；超过 0.5 秒的缺口不插值。Pearson 相关描述零滞后时的波形相似性；最大互相关描述时间对齐后的最强相似性；归一化均方根误差（NRMSE）用 Shimmer 信号标准差归一化。现场分析中，信号用加速度运动幅值和陀螺仪幅值概括，以降低传感器之间小方向差异的影响。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.1.2 Classification metrics",
        "分类中，正类是骨盆旋转。准确率报告所有窗口中分类正确的比例；平衡准确率平均两个类别的敏感性，因此较少受占多数的静态坐姿类别影响。精确率、召回率和 F1 分数针对骨盆旋转报告，因为这一类别决定无活动计时器是否被重置。这些指标只在留出参与者上计算。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.1.3 Co-located Agreement",
        "共位记录的对齐时间序列显示，XIAO 和 Shimmer 在所有加速度和陀螺仪通道上具有相同的主要运动峰值。所有轴上，加速度相关系数为 0.856-0.923，陀螺仪相关系数为 0.816-0.864。估计滞后保持在 -0.020 到 0 秒之间，NRMSE 为 0.43-0.61。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.1.4 Long Field-Worn Trend Agreement",
        "现场佩戴记录中，加速度运动幅值在 P1 和 P2 上的相关分别为 0.918 和 0.855，陀螺仪幅值相关分别为 0.908 和 0.864。滞后校正后，最大互相关达到 0.915-0.981。估计最佳滞后为 1.2-1.3 秒。后续检测分析只使用 XIAO 数据，并把目标定义为是否存在骨盆运动，而不是解剖学骨盆角度估计。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.2.1 Participants and Dataset",
        "五名志愿者在骶骨附近佩戴原型。研究通过 KU Leuven PRET 委员会批准，编号为 G-2025-10284-R2(MIN)，项目名称为“Daily Monitoring of Pelvic Rotation Using Wearable IMUs”，日期为 2026 年 1 月 5 日；所有参与者都提供知情同意。本论文没有分析人口统计或临床变量，因此初步结果不能按年龄、性别、身体组成、坐姿习惯或下背痛史分层。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.2.1 Dataset construction",
        "分类分析只使用列车旅途中带标签的坐姿片段。参与者被要求在静态坐姿和有意前后向骨盆旋转之间交替。路线中的非坐姿部分、无标签行和带电池断连缺口的片段从分类器数据集中排除。最终得到 98 个干净带标签片段文件，其中 66 个静态坐姿片段，32 个骨盆旋转片段。参与者 1-4 用于训练和内部模型检查，参与者 5 作为最终测试的留出对象。参与者 5 不用于特征选择、缩放、KNN 的 k 值选择或原型数量选择。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.2.1 Window counts",
        "窗口数量如下：P1 有 2,425 个静态窗口和 102 个骨盆旋转窗口；P2 有 2,641 个静态窗口和 70 个骨盆旋转窗口；P3 有 1,912 个静态窗口和 106 个骨盆旋转窗口；P4 有 2,014 个静态窗口和 121 个骨盆旋转窗口；P5 有 3,497 个静态窗口和 165 个骨盆旋转窗口。所有窗口均为 2 秒窗口、1 秒步长。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.2.2 Classifier Comparison",
        "使用参与者 1-4 的同一 90 维特征矩阵训练五种传统机器学习分类器，并在参与者 5 上评估：逻辑回归、RBF 核支持向量机、决策树、500 棵树的随机森林和 KNN。划分是跨人独立的，因此留出参与者没有任何窗口用于训练。对于逻辑回归、SVM 和 KNN，标准化步骤位于 scikit-learn 管线内部，只在参与者 1-4 上拟合，然后应用到参与者 5。参与者 5 没有用于 KNN 模型选择；在参与者 1-4 上针对 k=3、5、7 的留一参与者验证中，骨盆旋转平均 F1 分数分别为 0.758、0.764 和 0.767，因此最终测试前选择 k=7。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.2.2 Classifier results",
        "留出测试集中静态坐姿与骨盆旋转约为 21:1，因此除准确率外还报告平衡准确率和骨盆旋转 F1。KNN 的 F1 最高，为 0.957。165 个骨盆旋转窗口中，155 个被正确识别，10 个漏检；3,497 个静态坐姿窗口中，3,493 个正确分类，4 个被误判为骨盆旋转。各模型结果为：LR 准确率 0.909、F1 0.474；SVM 准确率 0.964、F1 0.709；DT 准确率 0.978、F1 0.781；RF 准确率 0.993、F1 0.911；KNN 准确率 0.996、F1 0.957。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.2.3 Compact KNN performance",
        "紧凑原型 KNN（12 个特征、128 个 KMeans 原型）在留出参与者上达到准确率 0.996、平衡准确率 0.963、精确率 0.994、召回率 0.927、F1 分数 0.959。完整基线使用 90 个特征和 9,391 个参考窗口，而紧凑模型使用 12 个特征和 128 个原型。因此，紧凑模型的精确率略高、召回率略低；同时，这一比较改变了特征集和参考集大小。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.2.3 On-device timing",
        "在 XIAO nRF54L15 Sense 上，使用硬件周期计数器测量了 24 个连续非重叠 2 秒窗口的每窗口推理时间。测量包括紧凑特征提取路径和对存储原型的 KNN 距离计算。KNN 推理本身平均 2.70 ms，范围 2.67-2.73 ms；特征加 KNN 总计平均 2.73 ms，范围 2.70-2.76 ms。"
    ),
    (
        "Chapter 4: Evaluation",
        "4.2.3 Compute footprint and boundary",
        "2.73 ms 总时间低于 2,000 ms 窗口周期的 0.15%。这个百分比只描述计算时间，不是功耗或电池寿命测量。部署固件使用 56.3 KB flash 和 9.3 KB RAM，在目标设备可用 flash 和 RAM 预算内，为 Zephyr RTOS、BLE 栈和未来扩展留下较大余量。BLE 通信、IMU 轮询占空比、稳压器损耗、睡眠状态行为和触觉驱动能量没有被测量。"
    ),
    (
        "Chapter 5: Discussion",
        "5.1 Position of the System",
        "原型达到了论文的主要工程目标：同一个小型可穿戴设备可以在骨盆附近感知运动，在实验期间流式传输带标签数据，在本地运行分类器，并驱动振动电机。这使得本研究位于久坐行为干预和可穿戴运动感知之间。Take-a-Stand 和 Move More @ Work 等工作场所干预通过组织或定时提示改变坐姿模式；相比之下，本文关注的是相关骨盆运动能否由身体佩戴设备本地感知并触发响应。因此，本文不替代行为干预研究，而是提供一个以后可以放入此类研究中评估的技术组件。"
    ),
    (
        "Chapter 5: Discussion",
        "5.1 Comparison with camera and furniture systems",
        "身体佩戴设计也不同于摄像头和家具系统。视觉系统可以在准备好的空间中分类坐姿，智能坐垫可以测量某个座位上的压力模式。本文系统牺牲了完整姿态重建和座位级压力细节，以换取跨座位跟随用户的能力。这个取舍符合论文目标，因为目标不是临床姿势诊断，而是检测长时间坐姿中是否发生了骨盆运动。"
    ),
    (
        "Chapter 5: Discussion",
        "5.1 Integration as the main value",
        "原型的实际价值在于集成。相关工作中每个单独部件都已经存在：IMU 可以测量运动，手机可以采集数据，分类器可以区分活动类别，触觉电机可以传递提示。本文贡献是把这些部分连接到一条聚焦骨盆的管线中，并用真实现场记录进行测试。因此，该原型应作为可行性系统来评价，而不是作为成熟干预产品。"
    ),
    (
        "Chapter 5: Discussion",
        "5.2 Sensor Validity and Placement",
        "Shimmer 比较支持将 XIAO 数据用于本文检测任务的趋势级分析，但不应解读为解剖角度验证。共位波形一致性足以说明自制节点捕捉到了运动峰值的时间，而 NRMSE 表明两种设备产生的幅值并不完全相同。这与可穿戴脊柱文献一致：IMU 可用于运动评估，但传感器位置、校准和分析目标会显著影响解释。对本文而言，幅值差异不如事件时间关键，因为分类器完全在 XIAO 数据上训练和测试。"
    ),
    (
        "Chapter 5: Discussion",
        "5.2 Placement and Liao et al.",
        "骶骨/下背部放置遵循与 Liao 等人骨盆 IMU 工作相同的逻辑，即传感器位置是骨盆旋转检测的核心因素。差异在于范围：Liao 等人提供基准式数据集和位置比较，而本文增加了低成本原型、移动记录、触觉输出和嵌入式部署。代价是数据集更小、变化更少。1.2-1.3 秒的现场滞后还说明真实部署受记录链路影响，而不仅是传感器本身。由于共位滞后接近零，现场偏移最可能来自 BLE 传输和 Android 端时间戳，而不是惯性传感器。"
    ),
    (
        "Chapter 5: Discussion",
        "5.2 Strength of validity claim",
        "这个区分影响有效性主张的强度。Shimmer 结果使得将 XIAO 用于二元运动检测任务是合理的，但不能证明该可穿戴设备能估计临床骨盆运动学。更强的生物力学验证需要受控放置、更大的配对数据集，以及专门用于解剖角度测量的参考系统。对当前论文而言，问题更窄：低成本设备是否捕捉到足够运动信息以支持分类器。"
    ),
    (
        "Chapter 5: Discussion",
        "5.3 Detection Performance",
        "分类结果对带标签初步任务而言很强，但行为意义较窄。任务区分静态坐姿和按指令完成的前后向骨盆旋转。它比一般人体活动识别简单，后者通常区分许多日常活动；也不同于坐姿分类，后者目标可能是完整姿态类别而不是运动事件的存在。KNN 的高 F1 说明在采集到的特征空间中两个指令类别可分，但不说明所有自发坐姿运动都能被识别。"
    ),
    (
        "Chapter 5: Discussion",
        "5.3 Why F1 matters more than accuracy",
        "分类器比较也解释了为什么 F1 比准确率更有信息量。留出集高度不平衡，单纯预测多数类也能获得高准确率却检测不到骨盆旋转。KNN 在精确率和召回率之间取得最佳平衡。对提醒逻辑而言，假阳性会在用户没有移动时重置无活动计时器；假阴性会在真实骨盆运动后未能重置计时器。KNN 结果同时包含两类错误，但假阳性尤其少。SVM 召回率更高但精确率低得多，会使提醒不可靠；随机森林更保守，漏检更多骨盆旋转窗口。逻辑回归表现差，因为分离关系可能无法由 90 维特征空间中的单个线性边界充分表示。"
    ),
    (
        "Chapter 5: Discussion",
        "5.3 Naturalness and reminder-specific metric",
        "最大开放问题是自然性。参与者执行的是指令式骨盆倾斜，而日常坐姿运动可能更小、更不周期，并且会与偶然移位、伸手、用手机或列车运动混合。因此，分类器目前只是在列车坐姿片段中检测受控指令动作。未来数据集必须包含办公室、家庭和交通中的自发坐姿行为，才能把模型描述为日常使用鲁棒。\n\n此外，该分类器针对提醒系统，而不是一般活动识别基准。在提醒系统中，错误成本取决于计时器。误检骨盆旋转会取消本应发生的提醒；漏检真实运动会让系统在用户已经移动后仍提醒。因此，类别特定的精确率和召回率比全局准确率更重要。"
    ),
    (
        "Chapter 5: Discussion",
        "5.4 Embedded Deployment",
        "嵌入式结果重要，因为它把感知任务连接到 TinyML 约束。许多 HAR 研究使用深度学习或较大的离线模型，但可穿戴微控制器受到内存、计算时间和能耗限制。部署模型选择了更保守的路径：时域特征、KNN 和 KMeans 原型压缩。这与轻量加速度特征仍可用于实时活动识别的早期工作一致；低能耗 HAR 硬件研究也强调在可穿戴健康应用中紧凑计算的重要性。"
    ),
    (
        "Chapter 5: Discussion",
        "5.4 Prototype compression interpretation",
        "KNN 通常是内存负担较重的惰性学习器，因为它存储训练集。用 128 个原型替代 9,391 个窗口改变了这一属性，使该方法可在 XIAO nRF54L15 Sense 上运行。虽然紧凑模型的留出 F1 略高，但不应解释为严格优于完整 KNN，因为特征集和参考集同时改变。更稳妥的解释是，原型压缩在将存储参考集减少约 73 倍的同时保留了相关决策行为。留出 F1 来自离线 2 秒窗口、1 秒步长评估，而嵌入式计时使用部署的非重叠 2 秒窗口周期。2.73 ms 每窗口运行时间远低于 2 秒窗口周期，但未测量功耗。因此，计时可行性尚不能证明电池寿命可行性。"
    ),
    (
        "Chapter 5: Discussion",
        "5.4 Embedded engineering detail",
        "部署结果还显示，在小型嵌入式系统中，工程细节可能支配结果。算法本身很轻量，但 Zephyr 主线程栈仍需增加，以避免启动故障。这提醒我们，嵌入式可行性不只是模型大小问题；固件调度、驱动、缓冲区和日志也会决定纸面上看似很小的模型是否真的能在设备上运行。"
    ),
    (
        "Chapter 5: Discussion",
        "5.5 Feedback and Intervention Scope",
        "DRV2605L 让原型具有不依赖手机通知的反馈通道。这一点很重要，因为预期用例是在坐姿期间提供本地提示，而不是事后可视化。不过，系统的触觉部分目前只完成了技术演示。久坐中断研究评估提示是否改变行为，生物反馈研究考察用户是否感知、接受并响应反馈。本文尚未回答这些问题。当前 10 分钟无活动阈值和 1 秒振动脉冲是实现选择，而不是验证过的行为参数。"
    ),
    (
        "Chapter 5: Discussion",
        "5.5 Platform claim and personalisation",
        "因此，完整系统应被描述为感知、分类和提示传递平台。它能在初步数据中检测指令动作并触发本地提醒，但还不能声称减少久坐时间、改善姿势或改变健康结果。这些主张需要用户研究来测量提醒感知、反复使用接受度、实际运动反应和长期依从性。\n\n另一个设计问题是个性化。当前系统使用同一个紧凑分类器进行留出参与者测试，适合初步跨人测试。但实际中，骨盆运动幅度、固定松紧、衣物、体型和坐姿环境都会改变信号分布。短时校准可以在保持微控制器内存预算内的同时，让原型集或无活动阈值适配个人。"
    ),
    (
        "Chapter 5: Discussion",
        "5.5 Label quality",
        "研究还显示了标签质量的重要性。由于目标运动细微且发生在坐姿中，后续处理很大程度依赖现场会话中记录的标注边界。Android 应用把标签直接写入传感器日志，减少了歧义，但协议仍依赖参与者在预定时间执行指令动作。未来采集应提供更明确的开始/停止提示、检查标签时间，并在隐私和伦理允许的情况下，为部分校准会话加入短视频或观察记录。"
    ),
    (
        "Chapter 5: Discussion",
        "5.6 Limitations",
        "局限性应从四个层面理解：数据集、分析管线、硬件和行为范围。数据集层面，分类器只在少量参与者上训练，并在一个留出参与者上测试。分析层面，干净坐姿片段和紧凑特征是人工选择的，尽管最终留出参与者没有参与特征选择、缩放和模型选择。硬件层面，原型足以进行现场采集，但尚不适合无人监督的长期佩戴。行为层面，系统可以发出提示，但研究没有测试提示是否改变坐姿行为。"
    ),
    (
        "Chapter 5: Discussion",
        "5.6 Listed limitations",
        "具体局限包括：样本量小，只有五名参与者，其中一名用于测试，无法从单个未见参与者得出总体结论；没有人口统计分层，未分析年龄、性别、身体组成、坐姿习惯和下背痛史；骨盆旋转是指令式前后倾斜，不是自发坐姿行为；只标注了单一运动方向，没有标注侧向骨盆旋转和其他坐姿活动模式；干净片段和紧凑特征为人工选择，虽然最终参与者被排除在选择之外，但流程仍应视为初步；Shimmer 覆盖有限，五个配对记录中三个失败，可靠性结论只基于两名参与者；分类器数据都来自列车旅途坐姿，尚未测试办公室和家庭使用；电池接触缺口可能在部署中造成错误无活动检测；没有电池寿命测量，只测了设备端计时；没有用户研究，振动提醒虽已设备端触发，但未评估感知、接受度或行为反应。"
    ),
    (
        "Chapter 5: Discussion",
        "5.7 Future Work",
        "未来需要更大、更多样的队列，以支持多个留出受试者交叉验证，并研究年龄、BMI 和下背痛史等因素。标注协议应从指令式骨盆旋转扩展到自发坐姿运动、侧向旋转、前倾和重心转移。硬件需要更可靠的电池连接和身体固定方式，并比较骶骨、下背部和腰部位置以明确最佳放置。算法方面，短时用户特异记录可用于个性化校准，使原型集适配个人；部署固件还应进行功耗分析，以明确连续 IMU 采样和分类下的电池寿命。要把系统描述为行为干预而不仅是感知和分类平台，还需要专门的用户研究。"
    ),
    (
        "Chapter 5: Discussion",
        "5.7 Validation phases and reproducibility",
        "下一项研究还应分离技术验证和行为评估。技术研究应首先测试办公室、家庭和交通情境中的传感器可靠性、电池寿命、佩戴稳定性和分类器泛化。之后的用户研究再考察提醒时机、振动强度、烦扰感、依从性和实际运动反应。分阶段可以让证据更清晰：先证明系统可靠检测目标行为，再测试提示是否以有用方式改变行为。\n\n最后一个实际步骤是提高嵌入式管线的可复现性。训练脚本、选中特征、原型值、固件头文件生成和计时测试应保存在一个有文档的工作流中。这样当更多参与者数据可用时，更容易更新模型，并降低离线版本和嵌入式版本分叉的风险。未来报告还应说明确切固件版本、手机型号、Android 版本和 BLE 记录设置，因为本文观察到的现场滞后说明数据路径会影响计时结果。"
    ),
    (
        "Chapter 6: Conclusion",
        "Conclusion opening",
        "本论文构建并测试了一个小型可穿戴系统，用于检测长时间坐姿中的骨盆运动，并在该运动缺失时提供触觉反馈。最终原型结合了佩戴在骶骨附近的 XIAO nRF54L15 Sense 传感节点、由 DRV2605L 驱动的硬币式振动电机、BLE 数据流、用于带标签现场记录的 Android 应用以及用于研究者侧监控的 PC 工具。评估分两阶段进行：先与 Shimmer 参考传感器比较，再开展五名参与者佩戴设备完成真实城市路线的初步现场研究。列车旅途中带标签坐姿数据用于训练和测试静态坐姿与骨盆旋转二分类器，最终紧凑模型部署到目标微控制器上。"
    ),
    (
        "Chapter 6: Conclusion",
        "Technical chain and claim boundary",
        "因此，论文展示了本地可穿戴提醒所需的完整技术链：感知、记录、标注、与参考传感器的趋势级比较、特征提取、分类器比较、原型压缩、嵌入式推理和触觉驱动。这比离线分类实验更完整，但比行为干预试验更窄。这个边界本身就是贡献的一部分：原型显示骨盆佩戴设备可以采集可用数据、本地分类目标运动，并在不依赖云处理的情况下提供触觉提示；但它尚未显示这些提示能改变长期坐姿行为、改善舒适度或降低健康风险。分离这些主张可以避免夸大实际效果，同时保留技术发现的价值。"
    ),
    (
        "Chapter 6: Conclusion",
        "Answer to RQ1",
        "对于 RQ1，嵌入式 IMU 对本检测任务而言能够捕捉到与研究级 Shimmer 参考设备足够一致的运动模式：共位 Pearson 相关系数为 0.816-0.923，滞后在 20 ms 内；现场最大互相关在滞后校正后最高达到 0.981。这是技术验证，不是临床验证，并且五名参与者中只有两名产生了有效配对 Shimmer 记录。"
    ),
    (
        "Chapter 6: Conclusion",
        "Answer to RQ2",
        "对于 RQ2，可穿戴平台支持真实交通情境中的实时数据采集和标注：五名参与者都在佩戴设备的情况下完成了现场路线，Android 应用用于会话期间活动标注。薄弱点主要是硬件可靠性，而不是软件工作流。"
    ),
    (
        "Chapter 6: Conclusion",
        "Answer to RQ3",
        "对于 RQ3，在跨人划分中（P1-P4 训练，P5 测试），k=7 的 KNN 分类器在骨盆旋转上达到 F1=0.957，165 个骨盆旋转窗口中检测到 155 个，3,497 个静态坐姿窗口中只有 4 个被误判。12 特征、128 原型的紧凑版本在同一留出参与者上达到 F1=0.959，并在 XIAO nRF54L15 Sense 上每 2 秒窗口运行 2.73 ms，相当于低于窗口周期 0.15% 的计算时间，占用 56.3 KB flash 和 9.3 KB RAM。因此，该检测器兼容目标设备本地部署，并足以支持这个初步指令式运动数据集。"
    ),
    (
        "Chapter 6: Conclusion",
        "Final conclusion",
        "本研究应被理解为可行性和初步验证研究，而不是行为干预试验。它提供了一条可运行的原型管线，结合了可穿戴硬件、移动与 PC 软件、现场数据、Shimmer 比较、可部署分类器和设备端计时。主要未解决问题包括：是否能推广到指令式骨盆旋转之外的自然运动，长时间佩戴下的机械可靠性，连续运行的电池寿命，以及用户对反复触觉提醒的反应。未来可以在这条管线基础上开展更多参与者、更自然坐姿运动、更可靠硬件和专门触觉提醒用户研究。总体而言，论文表明局部骨盆运动检测和提示传递在紧凑可穿戴平台上具有技术可行性；它的主要价值是把打断长时间静态坐姿的概念转化为可测量的原型，并使剩余弱点明确到足以指导下一轮迭代。"
    ),
    (
        "Appendices",
        "Appendix: IMU Output Data Rate Configuration",
        "XIAO nRF54L15 Sense 文档表明板载六轴 IMU 是 STMicroelectronics LSM6DS3TR-C。固件中通过 Zephyr 传感器 API 将加速度计和陀螺仪采样频率都请求为 208 Hz。LSM6DS3TR-C 数据表列出 208 Hz 是加速度计和陀螺仪均支持的输出数据率。干净日志片段中测得的有效读取率低于这个配置输出数据率，因为固件轮询传感器，同时执行 I2C、BLE 和操作系统任务。"
    ),
    (
        "Appendices",
        "Appendix: Coin Vibration Motor Specification",
        "原型使用的振动电机为 8 mm × 3 mm 硬币式迷你振动电机，供应商 listing 标称电压为 3-5 V、额定电流为 67 mA。由于未找到该具体零售件的正式数据表，论文还查阅了 Precision Microdrives 的相近 8 mm 硬币振动电机数据表作为规格参考。附录中复现了供应商 listing 的规格图片。"
    ),
    (
        "Appendices",
        "Appendix: Ethical Approval",
        "初步现场研究由 KU Leuven Privacy and Ethics platform（PRET）以项目“Daily Monitoring of Pelvic Rotation Using Wearable IMUs”审查批准，批准号 G-2025-10284-R2(MIN)，日期为 2026 年 1 月 5 日。申请由数据保护官评估其 GDPR 合规性，并由 Social and Societal Ethics Committee（SMEC）进行伦理标准审查；两者均给出有利决定。所有参与者在数据采集前提供知情同意，并可随时停止或退出。"
    ),
    (
        "Acknowledgements",
        "Acknowledgements",
        "致谢部分感谢导师 Prof. Bart Vanrumste 和共同导师 Prof. Jean-Marie Aerts、Meixing Liao 在整个论文期间提供指导、技术建议和反馈。他们在原型设计、骨盆旋转方法和写作方面的持续意见塑造了本研究。作者还感谢 Jona 在原型开发阶段提供技术支持，感谢参加初步现场研究并使数据集成为可能的志愿者，最后感谢家人和朋友在项目期间的耐心与支持。"
    ),
    (
        "Acronyms",
        "Acronyms",
        "缩略语：BLE = Bluetooth Low Energy，低功耗蓝牙；CSV = Comma-Separated Values，逗号分隔值；DT = Decision Tree，决策树；GDPR = General Data Protection Regulation，通用数据保护条例；GPS = Global Positioning System，全球定位系统；HAR = Human Activity Recognition，人体活动识别；I2C = Inter-Integrated Circuit，集成电路间总线；IMU = Inertial Measurement Unit，惯性测量单元；KNN = k-Nearest Neighbours，K 近邻；LOPO = Leave-One-Participant-Out，留一参与者外验证；LR = Logistic Regression，逻辑回归；MET = Metabolic Equivalent of Task，任务代谢当量；NRMSE = Normalised Root-Mean-Square Error，归一化均方根误差；PRET = Privacy and Ethics，隐私与伦理；RAM = Random-Access Memory，随机存取存储器；RBF = Radial Basis Function，径向基函数；RF = Random Forest，随机森林；RTOS = Real-Time Operating System，实时操作系统；SMEC = Social and Societal Ethics Committee，社会与社会伦理委员会；SVM = Support Vector Machine，支持向量机。"
    ),
    (
        "References",
        "Bibliography handling",
        "参考文献条目本身保留英文，不建议翻译到正式稿中。原因是论文题名、期刊名和出版信息需要保持可检索性和准确性。阅读时重点核对的是引用是否支持正文主张，而不是把文献列表翻译成中文。"
    ),
]


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.25)
    sec.bottom_margin = Cm(1.25)
    sec.left_margin = Cm(1.25)
    sec.right_margin = Cm(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("论文中文逐节全对照阅读版")
    set_font(r, size=18, bold=True)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run(
        "用于阅读、批注和逻辑核对；正式提交仍以英文排版稿为准。参考文献条目保持英文。"
    )
    set_font(r, size=10, color=(90, 90, 90))

    add_header_table(doc)
    current = None
    for section, left, right in CONTENT:
        if section != current:
            add_heading(doc, section)
            current = section
        add_pair(doc, left, right)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
