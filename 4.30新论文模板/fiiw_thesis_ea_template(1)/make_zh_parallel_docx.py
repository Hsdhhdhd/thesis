from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("thesis_中文全对照阅读版.docx")


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=9.5):
    cell.text = ""
    for idx, part in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(part)
        set_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Thesis Chinese-English Reading Companion\n论文中文全对照阅读版")
    set_font(r, size=18, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "说明：本文件用于阅读和核对，不替代正式排版稿。左栏为英文原文内容定位/摘要，右栏为中文对应译文。"
    )
    set_font(r, size=10, color=(90, 90, 90))


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, size=14, bold=True)


def add_pair(doc, left, right):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(12.7)
    table.columns[1].width = Cm(13.8)
    row = table.rows[0]
    row.cells[0].width = Cm(12.7)
    row.cells[1].width = Cm(13.8)
    set_cell_text(row.cells[0], left, size=9)
    set_cell_text(row.cells[1], right, size=9.5)


def add_header_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for c, text in zip(table.rows[0].cells, ["English source / 原文定位", "Chinese reading version / 中文对照"]):
        shade(c, "D9EAF7")
        set_cell_text(c, text, bold=True, size=10)


CONTENT = [
    (
        "Front matter: Title / Abstract",
        "Thesis title and abstract",
        "题目：一个面向肌萎缩侧索硬化症患者的轻量级可穿戴式手势识别界面。\n\n摘要：本文提出并评估一种轻量、低成本的可穿戴手势识别手套，用于帮助 ALS 患者在手部功能退化后仍能进行简单、可靠的数字交互。系统采用五个手指弯曲传感器和一个 MPU-6050 惯性测量单元，安装在柔性氨纶手套上；通过蓝牙把数据传输到计算机端，并用 Python 图形界面实时显示传感器状态、分类结果和系统延迟。\n\n研究聚焦两个问题：第一，使用者是否能够通过少量校准样本完成个性化手势识别；第二，一个简单、可解释、低算力的分类器是否足以达到可用性能。系统在三名 ALS 参与者和两名健康对照者上进行了评估，包含个体内交叉验证、留一参与者外验证、原型压缩、延迟测量和初步可用性反馈。\n\n结果显示，Full KNN 在多数参与者上准确率很高；压缩原型 KNN 在保持相近准确率的同时显著减少了模型规模。端到端推理时间保持在毫秒级，远低于交互系统常见的延迟阈值。研究表明，基于弯曲传感器和惯性传感器的个性化手势界面具备可行性，但仍需更多 ALS 用户、长期家庭场景测试以及更严格的疲劳和电池寿命评估。"
    ),
    (
        "Chapter 1: Introduction - clinical and technical motivation",
        "ALS is a progressive neurodegenerative disease. It affects motor neurons, causes weakness and loss of voluntary movement, and may eventually impair speech and communication. Assistive technologies can preserve autonomy, but many systems are expensive, bulky, visually demanding, or insufficiently adapted to residual hand motion.",
        "第 1 章引言首先建立临床和工程背景。ALS 是一种渐进性神经退行性疾病，会损伤运动神经元，引发肌肉无力、萎缩、痉挛以及自主运动能力下降。随着疾病进展，患者可能逐渐丧失稳定操作键盘、鼠标或触屏的能力，部分患者还会出现言语障碍。因此，辅助交互系统的目标不是追求复杂功能，而是在患者仍保留少量可控运动时，把这些运动转化为稳定、低负担的输入方式。\n\n现有辅助技术包括眼动追踪、脑机接口、肌电输入、语音控制和可穿戴传感器。它们各有优势，但也存在成本、佩戴复杂度、疲劳、鲁棒性、隐私和日常维护方面的问题。本文选择手套式可穿戴方案，是因为它能够直接捕捉残余手部动作，硬件简单，用户学习成本低，并且可以在不依赖视觉注意力的情况下完成若干基础命令。"
    ),
    (
        "Chapter 1: Introduction - literature gap and objective",
        "Prior work has shown that wearable gloves, flexible sensors, IMUs, EMG, and machine learning can recognize gestures. However, many studies use healthy participants, large gesture vocabularies, offline settings, or complex models. Few evaluate lightweight personalized systems with ALS users under real-time constraints.",
        "引言中的核心研究空白是：已有手势识别研究证明了传感器和机器学习的可行性，但很多工作主要在健康被试上完成，手势集合偏大，离线评估较多，或者依赖复杂模型；这与 ALS 患者的实际需求并不完全一致。ALS 用户更需要的是少量可重复动作、快速个体校准、低延迟反馈、低成本硬件和可解释算法。\n\n因此，本文目标是构建并验证一个轻量级、个性化、实时的手势识别原型。它不是要替代所有高端辅助沟通设备，而是探索一种低门槛的补充式交互方式：当患者还能完成有限手指弯曲或手腕姿态变化时，系统能否可靠识别这些动作，并把它们映射到简单的数字命令。"
    ),
    (
        "Chapter 1: Contributions",
        "The thesis contributes a glove prototype, a real-time software pipeline, a calibration workflow, a KNN-based classification baseline, evaluation with ALS participants and controls, and a discussion of limitations for deployment.",
        "本文贡献可概括为六点：第一，设计并制作了一套由五个弯曲传感器和一个 MPU-6050 IMU 组成的低成本手套原型；第二，开发了从 Arduino 采集、蓝牙传输、Python 端解析到 GUI 实时显示的完整软件链路；第三，设计了适合个体差异的校准流程，使每名用户可用少量样本建立个人模型；第四，实现并比较 Full KNN 与压缩原型 KNN，以测试轻量化模型是否能保持识别性能；第五，在 ALS 参与者和健康对照者上进行了初步实验，包括准确率、平衡准确率、精确率、召回率、F1 和延迟；第六，从辅助技术部署角度讨论了系统的局限，包括样本规模、真实家庭环境、电池寿命、长期疲劳和临床异质性。"
    ),
    (
        "Chapter 2: System Design - design requirements",
        "The design emphasizes low cost, comfort, repeatability, real-time operation, and personalization. The system should work with a small set of gestures and should not require high computational resources.",
        "第 2 章系统设计围绕辅助交互的实际约束展开。系统必须足够轻，不显著增加手部负担；硬件应便宜、易采购、易维修；传感器布局应尽量稳定，以保证同一手势在不同采样时具有可重复信号；算法应能在普通笔记本电脑上实时运行；校准流程应短，以降低患者疲劳。\n\n设计上没有追求识别大量复杂手势，而是采用少量、功能明确、可重复的手势类别。这个取舍符合 ALS 场景：稳定性、低疲劳和个体适配比手势数量更重要。"
    ),
    (
        "Chapter 2: Hardware architecture",
        "Five flex sensors are routed along the fingers. An MPU-6050 inertial unit measures acceleration and angular velocity. Arduino Nano reads analog and I2C signals and sends packets over Bluetooth through HC-05. The glove is made from stretch fabric.",
        "硬件架构包括四部分。第一，五个弯曲传感器分别沿五根手指布置，用于测量手指弯曲程度；第二，MPU-6050 惯性测量单元提供加速度和角速度信息，用来补充手腕姿态和整体运动特征；第三，Arduino Nano 负责采集模拟信号和 I2C 数据，并打包成串口数据；第四，HC-05 蓝牙模块把数据无线传输到电脑端。\n\n手套主体采用柔性氨纶材料，目的是降低穿戴阻碍。传感器固定方式需要在稳定性和舒适性之间折中：固定过松会造成漂移，固定过紧会增加疲劳或限制运动。论文中这部分说明了为什么选择这种硬件组合，也解释了它相对高端传感系统的优势：便宜、低功耗、结构透明、便于复现。"
    ),
    (
        "Chapter 2: Gesture vocabulary and interaction mapping",
        "The gesture set includes relaxed and finger or orientation-based gestures. These gestures are intended as simple command inputs rather than detailed natural hand pose reconstruction.",
        "手势集合设计为面向命令输入，而不是完整重建手部姿态。论文中的手势包括静息/放松状态、若干手指弯曲动作以及由 IMU 支持区分的姿态动作。每个手势都应满足三个条件：用户能相对轻松地做出；传感器信号与其他手势有可分辨差异；动作在多次重复中具有稳定性。\n\n这种设计逻辑与疾病场景一致。对 ALS 患者来说，过多、过细或需要快速切换的手势会增加疲劳并降低可靠性。因此论文强调的是小规模手势库和个体化校准，而不是通用大模型。"
    ),
    (
        "Chapter 3: Implementation - firmware and packet format",
        "Arduino firmware samples flex sensors and IMU data, formats them into a consistent packet, and transmits them over Bluetooth. A stable packet format is necessary for real-time parsing and classification.",
        "第 3 章实现部分说明了固件和上位机软件如何协同工作。Arduino 端周期性读取五个弯曲传感器和 IMU 数据，并把它们整理为固定格式的数据包发送。固定数据格式很重要，因为 Python 端需要持续解析串口流，如果字段顺序或分隔不稳定，会导致丢包、错位或错误分类。\n\n这一部分的重点不是复杂嵌入式控制，而是建立一条稳定的数据链：采样、格式化、传输、解析、特征化、分类和反馈。"
    ),
    (
        "Chapter 3: Python GUI and real-time pipeline",
        "The Python application provides a GUI for live sensor display, calibration, model training, prediction, and logging. The software keeps the pipeline simple enough to inspect and debug.",
        "Python 端程序承担实时交互任务。它连接蓝牙串口，读取数据包，显示当前传感器值，支持用户采集各类手势样本，训练分类器，并在运行阶段输出预测手势和响应时间。GUI 的作用不仅是展示结果，也让研究者可以检查传感器是否工作正常、某一通道是否漂移、样本是否采集错误。\n\n实现策略偏向可解释和可调试，而不是黑箱式自动化。这对原型研究很重要，因为在 ALS 用户身上采集数据时，错误可能来自传感器松动、动作疲劳、个体运动范围变化、蓝牙连接不稳或分类器本身。清晰的 GUI 有助于定位问题。"
    ),
    (
        "Chapter 3: Classification method",
        "A K-nearest-neighbor classifier is used because it is simple, non-parametric, personalized, and can work with limited calibration data. A compact prototype variant reduces memory and computation by representing each class with selected prototypes.",
        "分类方法采用 K 近邻。选择 KNN 的理由是它简单、可解释、适合小样本个体化校准，不需要长时间训练，也不需要复杂超参数。对于每名用户，系统用其个人校准样本建立模型，这比直接使用跨用户通用模型更适合 ALS 场景，因为患者手部运动能力差异很大。\n\n论文同时实现了压缩原型 KNN。Full KNN 保存全部训练样本，预测时需要与所有样本比较；Compact prototype KNN 则为每个类别保留代表性原型，以降低存储和计算量。这样可以测试一个关键工程问题：在资源受限或未来嵌入式部署时，是否能用更小模型保持接近的识别效果。"
    ),
    (
        "Chapter 3: Latency and reliability considerations",
        "Latency is measured across the software inference path. The study focuses on real-time feasibility, packet processing, and classification response rather than full battery-life validation.",
        "实现章节还涉及延迟和可靠性。论文评估的是从接收到数据到完成分类输出的推理路径延迟，重点判断系统是否满足实时交互的基本要求。结果中的毫秒级延迟说明 KNN 本身和 Python 管线不会成为主要瓶颈。\n\n需要注意的是，论文没有真实电池寿命实验，因此相关表述应保持谨慎。可以说系统硬件具有低功耗潜力，或基于元件规格推测适合移动部署，但不能把未测量的电池寿命写成实验证据。"
    ),
    (
        "Chapter 4: Evaluation - participants and protocol",
        "The evaluation includes three ALS participants and two healthy controls. Participants performed the defined gestures during calibration and testing. Metrics include accuracy, balanced accuracy, precision, recall, F1, prototype count, and latency.",
        "第 4 章评估部分是论文证据的核心。实验对象包括三名 ALS 参与者和两名健康对照者。每名参与者按照预设手势集合进行校准和测试，系统记录传感器数据与标签，并在后续评估中计算识别性能。\n\n指标包括准确率、平衡准确率、精确率、召回率和 F1 值。由于类别数量不多且样本可能不完全均衡，平衡准确率和召回率比单纯准确率更能说明各类别是否都被稳定识别。论文还报告压缩模型保留的原型数量和推理延迟，用于体现部署效率。"
    ),
    (
        "Chapter 4: Within-participant performance",
        "Within-participant cross-validation reports high scores for most users. This supports the idea that individualized calibration can capture stable gesture patterns.",
        "个体内交叉验证结果整体较高，说明同一用户经过个人校准后，系统能够学习其手势信号模式。这是论文最有力的结果之一，因为 ALS 患者之间差异很大，个体化模型比跨用户泛化更现实。\n\n阅读时应注意，个体内高准确率不等于临床可部署。它主要证明在受控实验、短时间采集和相对规范的动作下，传感器与 KNN 管线可以识别手势。长期使用时，疲劳、佩戴位置变化、手部状态波动和环境干扰仍可能降低性能。"
    ),
    (
        "Chapter 4: Held-out participant evaluation",
        "Held-out participant 5 evaluation compares Full KNN and Compact prototype KNN. The compact model keeps performance close to the full model while reducing the stored prototypes.",
        "留一参与者外评估用于观察模型在未见用户上的表现。论文特别展示了第 5 名参与者的 held-out 结果，并比较 Full KNN 与 Compact prototype KNN。图中可见两种方法的准确率和 F1 接近，压缩模型在部分指标上甚至略高或相近，但召回率存在下降空间。\n\n这个结果支持压缩原型方法的工程可行性：减少样本数量并不必然导致明显性能损失。不过，由于参与者数量很少，不能把它解释为强泛化结论。更准确的表述是：初步结果显示，原型压缩在本数据集上可以保持接近 Full KNN 的性能，值得在更大样本中继续验证。"
    ),
    (
        "Chapter 4: Latency results",
        "The measured classification latency is far below typical interactive thresholds. Therefore, the simple KNN-based pipeline appears suitable for real-time feedback.",
        "延迟结果表明，系统分类响应处于毫秒级，远低于常见人机交互中可感知延迟阈值。这个结果说明当前算法和软件管线足以支持实时反馈，用户做出手势后系统可以迅速更新预测。\n\n但延迟结果应限于已测路径：它证明推理和上位机处理足够快，不等于完整无线系统在所有环境下都无延迟问题。实际部署还需要考虑蓝牙连接稳定性、系统启动时间、误触发恢复、长时间运行内存稳定性等因素。"
    ),
    (
        "Chapter 4: Usability feedback",
        "Participants provided preliminary feedback about comfort, learnability, and gesture effort. These comments are qualitative and should be treated as exploratory evidence.",
        "可用性反馈属于探索性证据。参与者对舒适度、学习难度、动作负担和系统反馈进行了初步评价。这部分能帮助解释为什么某些手势更稳定，或者为什么某些用户在特定类别上表现较差。\n\n由于样本量小、测试时间短、没有长期家庭使用，不能把这些反馈写成确定性可用性结论。更合适的说法是：初步反馈显示该方案具有可接受性，但舒适度、疲劳、穿戴稳定性和长期依从性仍需进一步研究。"
    ),
    (
        "Chapter 5: Discussion - interpretation of results",
        "The discussion separates objective results from interpretation. It explains why the system worked well, where it failed, and how this compares with broader literature on wearable gesture recognition and assistive technology.",
        "第 5 章讨论需要完成结果之外的解释。结果章节应主要客观展示数据；讨论章节则要回答为什么会这样、这个结果与其他研究有什么关系、它说明了什么、没有说明什么。\n\n论文当前讨论的主线可以这样理解：高个体内性能说明个体化校准有效；压缩原型保持相近性能说明轻量部署具有潜力；毫秒级延迟说明算法不是实时交互瓶颈；但小样本、短期测试、缺少真实电池寿命和家庭场景验证限制了结论外推。"
    ),
    (
        "Chapter 5: Comparison with prior work",
        "The thesis should compare vertically and horizontally: against wearable gloves, IMU-based recognition, EMG/BCI assistive systems, eye-tracking, and other ALS communication tools.",
        "讨论中的横向和纵向比较是老师特别指出的重点。横向比较是把本文与同类可穿戴手势识别、IMU 手势识别、柔性传感器手套、肌电输入、眼动系统和脑机接口进行比较。本文的优势在于硬件便宜、算法简单、校准直接、实时性强；劣势在于适用对象受限于仍有残余手部运动的患者，且目前证据规模较小。\n\n纵向比较是把本文放在 ALS 辅助技术发展链条中看。高端 AAC、眼动和 BCI 可以服务更严重阶段的患者，但成本、训练负担或环境要求更高；本文方案更适合作为早期或中期残余手部运动仍可利用时的补充输入方式。这样讨论能避免把本系统说成万能解决方案，也能更准确地体现它的应用位置。"
    ),
    (
        "Chapter 5: Limitations",
        "Main limitations include small sample size, participant heterogeneity, short evaluation duration, no real battery-life experiment, no long-term home deployment, limited gesture vocabulary, and possible sensor placement variation.",
        "局限性应明确列出并与结论边界对应。第一，样本量很小，三名 ALS 参与者不足以代表疾病不同阶段和不同运动能力。第二，ALS 个体差异很大，因此高个体内性能不能直接推广到所有用户。第三，实验时间短，没有覆盖长期疲劳、学习效应或病情变化。第四，没有真实电池寿命测试，因此不能把续航作为实验证据。第五，没有家庭环境部署，无法验证日常穿戴、蓝牙干扰、照护者协助和误操作恢复。第六，手势词汇有限，尚未证明复杂交互任务中的效率。\n\n这些局限不削弱原型研究价值，但要求论文结论保持审慎。"
    ),
    (
        "Chapter 5: Implications and future work",
        "Future work should recruit more ALS participants, test longitudinal home use, measure battery life, refine gesture selection, improve mounting comfort, and compare classifiers under matched conditions.",
        "未来工作可以分为临床、硬件和算法三个方向。临床上，应扩大 ALS 参与者数量，覆盖不同疾病阶段，并进行多日或多周家庭试验。硬件上，应改进传感器固定方式、线缆管理、手套尺寸适配和供电模块，并真实测量电池寿命。算法上，应在相同数据划分下比较 KNN、SVM、随机森林、轻量神经网络和自适应校准方法，同时加入漂移检测和在线更新机制。\n\n更进一步，可以把手势识别结果映射到真实 AAC 软件、智能家居控制或计算机快捷命令中，评估任务完成时间、错误恢复和用户主观负担。"
    ),
    (
        "Chapter 6: Conclusion",
        "The thesis concludes that a low-cost wearable glove with personalized KNN classification can recognize a small set of gestures with high preliminary accuracy and low latency, but larger and longer studies are required before clinical claims.",
        "结论应保持清楚而克制：本文证明了一种低成本手套式传感系统结合个体化 KNN 分类，在小样本 ALS 与健康对照实验中可以实现少量手势的高准确率识别，并且推理延迟足以支持实时反馈。压缩原型模型显示出减少存储和计算需求的潜力。\n\n同时，结论不能过度宣称临床有效性。当前证据支持“可行性”和“原型验证”，不支持“大规模临床部署已经成熟”。后续需要更大样本、更长时间、真实家庭环境、真实电池寿命和更完整可用性评估。"
    ),
    (
        "Appendix: Technical appendix and datasheet",
        "Appendices document source code, calibration details, hardware references, and the comparable vibration motor datasheet used for specification support.",
        "附录用于支撑正文中不宜展开的技术细节，包括程序代码、校准流程、硬件连接、传感器说明和元件资料。关于振动电机，精确的 Sourcing Map 数据表未找到，因此论文使用 Precision Microdrives 308-100 作为相近规格的参考资料。正文应明确这是 comparable datasheet，而不是声称找到了完全相同型号的官方数据表。\n\n附录材料的作用是提高可复现性，让读者能理解硬件选择、代码结构和实验流程。"
    ),
    (
        "Acknowledgements",
        "Acknowledgements thank the supervisor, participants, colleagues, and support network.",
        "致谢部分感谢导师、参与者、同学/同事以及提供支持的人。它不属于技术论证部分，但在最终稿中应保持语气正式、简洁，避免过于口语化。"
    ),
    (
        "List of symbols / acronyms",
        "ALS, AAC, IMU, KNN, GUI, F1, BLE/Bluetooth-related terms, and other abbreviations are defined.",
        "符号和缩略语列表用于统一术语。关键术语包括 ALS（肌萎缩侧索硬化症）、AAC（辅助与替代沟通）、IMU（惯性测量单元）、KNN（K 近邻）、GUI（图形用户界面）、F1 分数、蓝牙相关术语等。阅读正文时应保持这些译名一致，避免同一概念在不同章节出现不同中文说法。"
    ),
    (
        "Figures and tables",
        "Figures and tables include system architecture, glove prototype, GUI, performance charts, latency, participant results, and hardware/specification summaries.",
        "图表主要承担两类作用：一类说明系统结构，如手套硬件、数据流、GUI 和传感器布置；另一类呈现实验结果，如不同参与者的分类性能、Full KNN 与 Compact prototype KNN 比较、延迟统计和硬件规格。\n\n阅读时建议先看图表标题和坐标轴，再回到正文解释。结果图只说明“发生了什么”；为什么好、为什么差、与文献相比如何，应在讨论章节寻找。"
    ),
    (
        "References",
        "The bibliography is kept in English in the formal thesis. Citation appropriateness should be checked against source abstracts and claims.",
        "参考文献列表在正式论文中保留英文题名和出版信息，不建议翻译成中文后放入正式稿，以免造成题名不准确或检索困难。阅读对照时，重点应检查每个引用是否真正支持对应句子：如果正文说的是 ALS 流行病学，就应引用 ALS 流行病学资料；如果说的是手势识别算法，就应引用可穿戴识别或机器学习资料；如果说的是具体元件规格，就应引用数据表或厂家资料。"
    ),
]


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.3)
    sec.bottom_margin = Cm(1.3)
    sec.left_margin = Cm(1.3)
    sec.right_margin = Cm(1.3)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10)

    add_title(doc)
    add_header_table(doc)
    current = None
    for section, left, right in CONTENT:
        if section != current:
            add_section(doc, section)
            current = section
        add_pair(doc, left, right)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "生成说明：本阅读版按论文结构制作中文对应解释，便于快速通读、批注和核对逻辑；正式提交仍以英文排版稿为准。"
    )
    set_font(r, size=9, color=(100, 100, 100))
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
